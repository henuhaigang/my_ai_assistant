import os
import tempfile
from typing import Dict, Any
from fastapi import UploadFile, HTTPException

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


async def analyze_file(file: UploadFile) -> Dict[str, Any]:
    
    suffix = os.path.splitext(file.filename)[1] if file.filename else ".tmp"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        content = await file.read()
        tmp_file.write(content)
        temp_path = tmp_file.name
    
    try:
        
        extracted_text = ""
        metadata = {
            "filename": file.filename,
            "content_type": file.content_type or "",
            "size_bytes": len(content),
            "file_extension": suffix.lower(),
            "page_count": 0
        }
        
        if not content:
            raise HTTPException(
                status_code=400, 
                detail="Empty file uploaded"
            )
            
        # PDF 文件分析
        if temp_path.endswith('.pdf'):
            try:
                import PyPDF2
                
                with open(temp_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    
                    metadata.update({
                        "page_count": len(reader.pages),
                        "file_type": "PDF"
                    })
                    
                    # 从PDF中提取文本
                    for page in reader.pages:
                        extracted_text += page.extract_text()
                        
            except Exception as e:
                raise HTTPException(
                    status_code=500,
                    detail=f"Error processing PDF file: {str(e)}"
                )
                
        # Word 文档分析
        elif temp_path.endswith('.docx'):
            if DocxDocument is None:
                raise HTTPException(
                    status_code=400,
                    detail="python-docx library not installed. Install with: pip install python-docx"
                )
                
            try:
                doc = DocxDocument(temp_path)
                
                metadata.update({
                    "paragraph_count": len(doc.paragraphs),
                    "word_count": sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip()),
                    "file_type": "Word Document"
                })
                
                # 从Word文档中提取文本
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        extracted_text += paragraph.text + "\n"
                        
            except Exception as e:
                raise HTTPException(
                    status_code=500,
                    detail=f"Error processing Word document: {str(e)}"
                )
                
        # Excel 文件分析
        elif temp_path.endswith('.xlsx') or temp_path.endswith('.xls'):
            if openpyxl is None:
                raise HTTPException(
                    status_code=400,
                    detail="openpyxl library not installed. Install with: pip install openpyxl"
                )
            
            try:
                wb = openpyxl.load_workbook(temp_path)
                
                metadata.update({
                    "sheet_count": len(wb.sheetnames),
                    "file_type": f"Excel ({temp_path.endswith('.xlsx') and 'XLSX' or 'XLS'})"
                })
                
                # 从Excel文件中提取文本
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    for row in ws.iter_rows(values_only=True):
                        for cell in row:
                            if cell is not None:
                                extracted_text += str(cell) + " "
                        extracted_text += "\n"
                        
            except Exception as e:
                raise HTTPException(
                    status_code=500,
                    detail=f"Error processing Excel file: {str(e)}"
                )
                
        # 文本文件分析
        elif temp_path.endswith('.txt'):
            try:
                with open(temp_path, 'r', encoding='utf-8') as f:
                    extracted_text = f.read()
                    
                metadata.update({
                    "line_count": len(extracted_text.splitlines()),
                    "file_type": "Text File"
                })
                
            except UnicodeDecodeError:
                # 如果UTF-8失败，尝试其他编码
                try:
                    with open(temp_path, 'r', encoding='latin-1') as f:
                        extracted_text = f.read()
                        
                    metadata.update({
                        "line_count": len(extracted_text.splitlines()),
                        "file_type": "Text File (Latin)"
                    })
                except Exception as e:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Error reading text file: {str(e)}"
                    )
                    
        # 通用文本提取（尝试从二进制文件中获取可读内容）
        else:
            try:
                with open(temp_path, 'r', encoding='utf-8') as f:
                    extracted_text = f.read()
                
                metadata.update({
                    "line_count": len(extracted_text.splitlines()),
                    "file_type": "Text File"
                })
                
            except UnicodeDecodeError:
                # 如果UTF-8失败，尝试其他编码
                try:
                    with open(temp_path, 'r', encoding='latin-1') as f:
                        extracted_text = f.read()
                    
                    metadata.update({
                        "line_count": len(extracted_text.splitlines()),
                        "file_type": "Text File (Latin)"
                    })
                except Exception as e:
                    # 如果所有方法都失败，返回空文本
                    metadata.update({
                        "file_type": "Unknown Binary"
                    })
                    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"File analysis failed: {str(e)}"
        )
        
    finally:
        # 清理临时文件
        if os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception:
                pass  # 忽略清理失败
                
    return {
        "filename": file.filename or "unknown",
        "content_type": file.content_type or "",
        "size_bytes": len(content),
        "extracted_text": extracted_text[:10000] if extracted_text else "",  
        "metadata": metadata,
        "success": True
    }